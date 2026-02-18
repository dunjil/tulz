"""OCR (Optical Character Recognition) service - Refactored.

Provides text extraction from images and PDFs using Tesseract OCR.
Now with proper async handling, image preprocessing, and quality assessment.

SYSTEM REQUIREMENTS:
- Tesseract OCR must be installed:
  Ubuntu: sudo apt-get install tesseract-ocr tesseract-ocr-eng poppler-utils
  macOS: brew install tesseract poppler

SUPPORTED LANGUAGES (Pro tier):
- eng (English) - default
- fra (French), deu (German), spa (Spanish), ita (Italian)
- por (Portuguese), nld (Dutch), pol (Polish), rus (Russian)
- chi_sim (Chinese Simplified), jpn (Japanese), kor (Korean), ara (Arabic)
"""

import asyncio
import io
import os
import uuid
from concurrent.futures import ThreadPoolExecutor
from typing import Optional

from PIL import Image

from app.config import settings
from app.core.exceptions import BadRequestError
from app.services.tools.ocr_preprocessing import preprocess_image, assess_image_quality
from app.services.tools.ocr_quality import assess_ocr_quality, detect_language_from_text


# Supported languages for OCR
SUPPORTED_LANGUAGES = {
    "eng": "English",
    "fra": "French",
    "deu": "German",
    "spa": "Spanish",
    "ita": "Italian",
    "por": "Portuguese",
    "nld": "Dutch",
    "pol": "Polish",
    "rus": "Russian",
    "chi_sim": "Chinese (Simplified)",
    "chi_tra": "Chinese (Traditional)",
    "jpn": "Japanese",
    "kor": "Korean",
    "ara": "Arabic",
    "hin": "Hindi",
    "tha": "Thai",
    "vie": "Vietnamese",
    "tur": "Turkish",
}

# Free tier only gets English
FREE_TIER_LANGUAGES = {"eng"}

# Tier limits
TIER_LIMITS = {
    "free": {
        "max_file_size_mb": 2,
        "max_pdf_pages": 5,
        "languages": ["eng"],
        "timeout_seconds": 15,
        "preprocessing": False,  # No preprocessing for free tier
    },
    "pro": {
        "max_file_size_mb": 10,
        "max_pdf_pages": 20,
        "languages": list(SUPPORTED_LANGUAGES.keys()),
        "timeout_seconds": 60,
        "preprocessing": True,  # Preprocessing available for pro
    },
}


class OCRService:
    """Service for OCR operations using Tesseract.
    
    Now properly handles async operations and includes preprocessing.
    """

    def __init__(self):
        self.temp_dir = settings.temp_file_dir
        os.makedirs(self.temp_dir, exist_ok=True)
        self._executor = ThreadPoolExecutor(max_workers=4)

    async def extract_text_from_image(
        self,
        image_bytes: bytes,
        language: str = "eng",
        tier: str = "free",
        use_preprocessing: bool = True,
    ) -> dict:
        """Extract text from an image using Tesseract OCR.
        
        Args:
            image_bytes: Raw image bytes
            language: Tesseract language code (e.g., 'eng', 'fra')
            tier: User tier for limit enforcement
            use_preprocessing: Apply image preprocessing for better accuracy
            
        Returns:
            dict with extracted text, quality metrics, and metadata
        """
        limits = TIER_LIMITS.get(tier, TIER_LIMITS["free"])

        # Validate language
        if language not in limits["languages"]:
            if tier == "free":
                raise BadRequestError(
                    message=f"Language '{language}' requires Pro tier. Free tier only supports English."
                )
            raise BadRequestError(message=f"Unsupported language: {language}")

        # Check file size
        file_size_mb = len(image_bytes) / (1024 * 1024)
        if file_size_mb > limits["max_file_size_mb"]:
            raise BadRequestError(
                message=f"File too large. {tier.title()} tier limit: {limits['max_file_size_mb']}MB"
            )

        # Run OCR in thread pool to avoid blocking
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(
            self._executor,
            self._sync_extract_text_from_image,
            image_bytes,
            language,
            limits,
            use_preprocessing and limits["preprocessing"],
        )

        return result

    def _sync_extract_text_from_image(
        self,
        image_bytes: bytes,
        language: str,
        limits: dict,
        use_preprocessing: bool,
    ) -> dict:
        """Synchronous OCR processing (runs in thread pool).
        
        This is the actual blocking operation that runs in a separate thread.
        """
        import pytesseract

        try:
            # Open image
            image = Image.open(io.BytesIO(image_bytes))

            # Convert to RGB if necessary
            if image.mode in ("RGBA", "P"):
                image = image.convert("RGB")

            # Assess image quality before processing
            image_quality = assess_image_quality(image)

            # Preprocess if enabled
            if use_preprocessing:
                image = preprocess_image(image)

            # Perform OCR
            text = pytesseract.image_to_string(
                image,
                lang=language,
                timeout=limits["timeout_seconds"],
            )

            # Get confidence data
            data = pytesseract.image_to_data(
                image,
                lang=language,
                output_type=pytesseract.Output.DICT,
                timeout=limits["timeout_seconds"],
            )

            # Calculate average confidence (excluding -1 values)
            confidences = [c for c in data["conf"] if c != -1]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0

            # Count words and characters
            word_count = len([w for w in text.split() if w.strip()])
            char_count = len(text)

            # Assess OCR quality
            ocr_quality = assess_ocr_quality(
                confidence=avg_confidence,
                word_count=word_count,
                char_count=char_count,
                language=language,
            )

            # Detect language from text
            detected_lang = detect_language_from_text(text)

            return {
                "success": True,
                "text": text.strip(),
                "word_count": word_count,
                "char_count": char_count,
                "confidence": round(avg_confidence, 1),
                "language": language,
                "language_name": SUPPORTED_LANGUAGES.get(language, language),
                "detected_language": detected_lang,
                "preprocessing_applied": use_preprocessing,
                "quality": ocr_quality,
                "image_quality": image_quality,
            }

        except Exception as e:
            raise BadRequestError(message=f"OCR failed: {str(e)}")

    async def create_searchable_pdf(
        self,
        pdf_bytes: bytes,
        language: str = "eng",
        tier: str = "free",
    ) -> dict:
        """Convert a scanned PDF to a searchable PDF using OCRmyPDF.
        
        Args:
            pdf_bytes: Raw PDF bytes
            language: Tesseract language code
            tier: User tier for limit enforcement
            
        Returns:
            dict with output file path and metadata
        """
        from pypdf import PdfReader

        limits = TIER_LIMITS.get(tier, TIER_LIMITS["free"])

        # Validate language
        if language not in limits["languages"]:
            if tier == "free":
                raise BadRequestError(
                    message=f"Language '{language}' requires Pro tier. Free tier only supports English."
                )
            raise BadRequestError(message=f"Unsupported language: {language}")

        # Check file size
        file_size_mb = len(pdf_bytes) / (1024 * 1024)
        if file_size_mb > limits["max_file_size_mb"]:
            raise BadRequestError(
                message=f"File too large. {tier.title()} tier limit: {limits['max_file_size_mb']}MB"
            )

        # Check page count
        try:
            reader = PdfReader(io.BytesIO(pdf_bytes))
            page_count = len(reader.pages)
        except Exception:
            raise BadRequestError(message="Invalid PDF file")

        if page_count > limits["max_pdf_pages"]:
            raise BadRequestError(
                message=f"PDF has {page_count} pages. {tier.title()} tier limit: {limits['max_pdf_pages']} pages"
            )

        # Run OCRmyPDF in thread pool
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(
            self._executor,
            self._sync_create_searchable_pdf,
            pdf_bytes,
            language,
            page_count,
        )

        return result

    def _sync_create_searchable_pdf(
        self,
        pdf_bytes: bytes,
        language: str,
        page_count: int,
    ) -> dict:
        """Synchronous PDF OCR processing (runs in thread pool)."""
        import ocrmypdf

        try:
            # Create temp files
            input_path = os.path.join(self.temp_dir, f"ocr_input_{uuid.uuid4().hex}.pdf")
            output_path = os.path.join(self.temp_dir, f"ocr_output_{uuid.uuid4().hex}.pdf")

            # Write input PDF
            with open(input_path, "wb") as f:
                f.write(pdf_bytes)

            # Run OCRmyPDF
            ocrmypdf.ocr(
                input_path,
                output_path,
                language=language,
                skip_text=True,  # Skip pages that already have text
                optimize=1,  # Light optimization
                progress_bar=False,
                jobs=1,  # Single thread to limit CPU usage
            )

            # Get output file size
            output_size = os.path.getsize(output_path)

            # Cleanup input file
            os.remove(input_path)

            return {
                "success": True,
                "output_path": output_path,
                "filename": os.path.basename(output_path),
                "page_count": page_count,
                "size": output_size,
                "language": language,
                "language_name": SUPPORTED_LANGUAGES.get(language, language),
            }

        except ocrmypdf.exceptions.PriorOcrFoundError:
            # PDF already has text
            if os.path.exists(input_path):
                os.remove(input_path)
            raise BadRequestError(
                message="This PDF already contains searchable text. No OCR needed."
            )
        except Exception as e:
            # Cleanup on error
            if os.path.exists(input_path):
                os.remove(input_path)
            if os.path.exists(output_path):
                os.remove(output_path)
            raise BadRequestError(message=f"OCR failed: {str(e)}")

    async def extract_text_from_pdf(
        self,
        pdf_bytes: bytes,
        language: str = "eng",
        tier: str = "free",
        use_preprocessing: bool = True,
    ) -> dict:
        """Extract text from a scanned PDF.
        
        Args:
            pdf_bytes: Raw PDF bytes
            language: Tesseract language code
            tier: User tier for limit enforcement
            use_preprocessing: Apply image preprocessing
            
        Returns:
            dict with extracted text and metadata
        """
        from pypdf import PdfReader

        limits = TIER_LIMITS.get(tier, TIER_LIMITS["free"])

        # Validate language
        if language not in limits["languages"]:
            if tier == "free":
                raise BadRequestError(
                    message=f"Language '{language}' requires Pro tier. Free tier only supports English."
                )
            raise BadRequestError(message=f"Unsupported language: {language}")

        # Check file size
        file_size_mb = len(pdf_bytes) / (1024 * 1024)
        if file_size_mb > limits["max_file_size_mb"]:
            raise BadRequestError(
                message=f"File too large. {tier.title()} tier limit: {limits['max_file_size_mb']}MB"
            )

        # Check page count
        try:
            reader = PdfReader(io.BytesIO(pdf_bytes))
            page_count = len(reader.pages)
        except Exception:
            raise BadRequestError(message="Invalid PDF file")

        if page_count > limits["max_pdf_pages"]:
            raise BadRequestError(
                message=f"PDF has {page_count} pages. {tier.title()} tier limit: {limits['max_pdf_pages']} pages"
            )

        # Run PDF OCR in thread pool
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(
            self._executor,
            self._sync_extract_text_from_pdf,
            pdf_bytes,
            language,
            limits,
            use_preprocessing and limits["preprocessing"],
        )

        return result

    def _sync_extract_text_from_pdf(
        self,
        pdf_bytes: bytes,
        language: str,
        limits: dict,
        use_preprocessing: bool,
    ) -> dict:
        """Synchronous PDF text extraction (runs in thread pool)."""
        import pytesseract
        from pdf2image import convert_from_bytes

        try:
            # Convert PDF pages to images
            images = convert_from_bytes(
                pdf_bytes,
                dpi=200,  # Good balance of quality vs speed
                thread_count=1,  # Limit CPU usage
            )

            # Extract text from each page
            all_text = []
            total_confidence = 0
            confidence_count = 0

            for i, image in enumerate(images):
                # Preprocess if enabled
                if use_preprocessing:
                    image = preprocess_image(image)

                # OCR the page
                text = pytesseract.image_to_string(
                    image,
                    lang=language,
                    timeout=limits["timeout_seconds"],
                )
                all_text.append(f"--- Page {i + 1} ---\n{text.strip()}")

                # Get confidence
                data = pytesseract.image_to_data(
                    image,
                    lang=language,
                    output_type=pytesseract.Output.DICT,
                    timeout=limits["timeout_seconds"],
                )
                confidences = [c for c in data["conf"] if c != -1]
                if confidences:
                    total_confidence += sum(confidences)
                    confidence_count += len(confidences)

            combined_text = "\n\n".join(all_text)
            avg_confidence = total_confidence / confidence_count if confidence_count else 0
            word_count = len([w for w in combined_text.split() if w.strip()])
            char_count = len(combined_text)

            # Assess OCR quality
            ocr_quality = assess_ocr_quality(
                confidence=avg_confidence,
                word_count=word_count,
                char_count=char_count,
                language=language,
            )

            # Detect language
            detected_lang = detect_language_from_text(combined_text)

            return {
                "success": True,
                "text": combined_text,
                "word_count": word_count,
                "char_count": char_count,
                "page_count": len(images),
                "confidence": round(avg_confidence, 1),
                "language": language,
                "language_name": SUPPORTED_LANGUAGES.get(language, language),
                "detected_language": detected_lang,
                "preprocessing_applied": use_preprocessing,
                "quality": ocr_quality,
            }

        except Exception as e:
            raise BadRequestError(message=f"OCR failed: {str(e)}")

    async def ocr_to_word(
        self,
        file_bytes: bytes,
        file_type: str,  # 'image' or 'pdf'
        language: str = "eng",
        tier: str = "free",
        use_preprocessing: bool = True,
    ) -> dict:
        """Convert scanned image/PDF to Word document.
        
        Args:
            file_bytes: Raw file bytes
            file_type: 'image' or 'pdf'
            language: Tesseract language code
            tier: User tier for limit enforcement
            use_preprocessing: Apply preprocessing
            
        Returns:
            dict with output file path and metadata
        """
        # First extract text
        if file_type == "image":
            result = await self.extract_text_from_image(
                file_bytes, language, tier, use_preprocessing
            )
        else:
            result = await self.extract_text_from_pdf(
                file_bytes, language, tier, use_preprocessing
            )

        if not result["text"]:
            raise BadRequestError(message="No text detected in the file")

        # Create Word document in thread pool
        loop = asyncio.get_event_loop()
        word_result = await loop.run_in_executor(
            self._executor,
            self._sync_create_word_document,
            result,
            language,
        )

        return word_result

    def _sync_create_word_document(self, ocr_result: dict, language: str) -> dict:
        """Create Word document from OCR result (runs in thread pool)."""
        from docx import Document
        from docx.shared import Pt

        try:
            # Create Word document
            doc = Document()

            # Add title
            doc.add_heading("OCR Extracted Text", 0)

            # Add metadata
            doc.add_paragraph(
                f"Language: {ocr_result['language_name']} | "
                f"Words: {ocr_result['word_count']} | "
                f"Confidence: {ocr_result['confidence']}%"
            )
            
            # Add quality info if available
            if "quality" in ocr_result:
                quality_info = ocr_result["quality"]
                if quality_info.get("warnings"):
                    doc.add_paragraph("⚠️ Quality Warnings:")
                    for warning in quality_info["warnings"]:
                        doc.add_paragraph(f"  • {warning}", style="List Bullet")
            
            doc.add_paragraph("")  # Spacer

            # Add extracted text
            for paragraph in ocr_result["text"].split("\n\n"):
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())
                    p.style.font.size = Pt(11)

            # Save to temp file
            output_filename = f"ocr_document_{uuid.uuid4().hex[:8]}.docx"
            output_path = os.path.join(self.temp_dir, output_filename)
            doc.save(output_path)

            return {
                "success": True,
                "output_path": output_path,
                "filename": output_filename,
                "size": os.path.getsize(output_path),
                "word_count": ocr_result["word_count"],
                "confidence": ocr_result["confidence"],
                "language": language,
                "language_name": ocr_result["language_name"],
                "quality": ocr_result.get("quality"),
            }

        except Exception as e:
            raise BadRequestError(message=f"Failed to create Word document: {str(e)}")

    def get_supported_languages(self, tier: str = "free") -> list:
        """Get list of supported languages for a tier."""
        limits = TIER_LIMITS.get(tier, TIER_LIMITS["free"])
        return [
            {"code": lang, "name": SUPPORTED_LANGUAGES[lang]}
            for lang in limits["languages"]
            if lang in SUPPORTED_LANGUAGES
        ]

    def get_tier_limits(self, tier: str = "free") -> dict:
        """Get limits for a tier."""
        return TIER_LIMITS.get(tier, TIER_LIMITS["free"])
