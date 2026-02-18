"""OCR background tasks for Celery.

These tasks handle heavy OCR operations asynchronously.
"""

import os
import uuid
from typing import Optional

from celery import shared_task

from app.config import settings


@shared_task(bind=True, max_retries=2, default_retry_delay=30)
def extract_text_from_image_task(
    self,
    image_path: str,
    language: str = "eng",
    output_dir: str = None,
) -> dict:
    """Extract text from an image using Tesseract OCR.

    Args:
        image_path: Path to the image file
        language: Tesseract language code
        output_dir: Directory to save output files

    Returns:
        dict with extracted text and metadata
    """
    try:
        import pytesseract
        from PIL import Image

        output_dir = output_dir or settings.temp_file_dir

        # Open image
        image = Image.open(image_path)

        # Convert to RGB if necessary
        if image.mode in ("RGBA", "P"):
            image = image.convert("RGB")

        # Perform OCR
        text = pytesseract.image_to_string(image, lang=language, timeout=60)

        # Get confidence data
        data = pytesseract.image_to_data(
            image,
            lang=language,
            output_type=pytesseract.Output.DICT,
            timeout=60,
        )

        # Calculate average confidence
        confidences = [c for c in data["conf"] if c != -1]
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0

        # Save text to file
        text_filename = f"ocr_text_{uuid.uuid4().hex[:8]}.txt"
        text_path = os.path.join(output_dir, text_filename)
        with open(text_path, "w", encoding="utf-8") as f:
            f.write(text.strip())

        return {
            "success": True,
            "text": text.strip(),
            "word_count": len([w for w in text.split() if w.strip()]),
            "char_count": len(text),
            "confidence": round(avg_confidence, 1),
            "output_path": text_path,
            "filename": text_filename,
        }

    except Exception as e:
        self.retry(exc=e)


@shared_task(bind=True, max_retries=2, default_retry_delay=60)
def create_searchable_pdf_task(
    self,
    pdf_path: str,
    language: str = "eng",
    output_dir: str = None,
) -> dict:
    """Convert a scanned PDF to a searchable PDF using OCRmyPDF.

    Args:
        pdf_path: Path to the input PDF file
        language: Tesseract language code
        output_dir: Directory to save output file

    Returns:
        dict with output file path and metadata
    """
    try:
        import ocrmypdf
        from pypdf import PdfReader

        output_dir = output_dir or settings.temp_file_dir

        # Get page count
        reader = PdfReader(pdf_path)
        page_count = len(reader.pages)

        # Create output path
        output_filename = f"ocr_searchable_{uuid.uuid4().hex[:8]}.pdf"
        output_path = os.path.join(output_dir, output_filename)

        # Run OCRmyPDF
        ocrmypdf.ocr(
            pdf_path,
            output_path,
            language=language,
            skip_text=True,
            optimize=1,
            progress_bar=False,
            jobs=1,
        )

        return {
            "success": True,
            "output_path": output_path,
            "filename": output_filename,
            "page_count": page_count,
            "size": os.path.getsize(output_path),
        }

    except ocrmypdf.exceptions.PriorOcrFoundError:
        return {
            "success": False,
            "error": "PDF already contains searchable text",
        }
    except Exception as e:
        self.retry(exc=e)


@shared_task(bind=True, max_retries=2, default_retry_delay=60)
def extract_text_from_pdf_task(
    self,
    pdf_path: str,
    language: str = "eng",
    output_dir: str = None,
    max_pages: int = 20,
) -> dict:
    """Extract text from a scanned PDF.

    Args:
        pdf_path: Path to the input PDF file
        language: Tesseract language code
        output_dir: Directory to save output file
        max_pages: Maximum pages to process

    Returns:
        dict with extracted text and metadata
    """
    try:
        import pytesseract
        from pdf2image import convert_from_path
        from pypdf import PdfReader

        output_dir = output_dir or settings.temp_file_dir

        # Get page count
        reader = PdfReader(pdf_path)
        page_count = min(len(reader.pages), max_pages)

        # Convert PDF pages to images
        images = convert_from_path(
            pdf_path,
            dpi=200,
            thread_count=1,
            last_page=page_count,
        )

        # Extract text from each page
        all_text = []
        total_confidence = 0
        confidence_count = 0

        for i, image in enumerate(images):
            text = pytesseract.image_to_string(image, lang=language, timeout=60)
            all_text.append(f"--- Page {i + 1} ---\n{text.strip()}")

            # Get confidence
            data = pytesseract.image_to_data(
                image,
                lang=language,
                output_type=pytesseract.Output.DICT,
                timeout=60,
            )
            confidences = [c for c in data["conf"] if c != -1]
            if confidences:
                total_confidence += sum(confidences)
                confidence_count += len(confidences)

        combined_text = "\n\n".join(all_text)
        avg_confidence = total_confidence / confidence_count if confidence_count else 0

        # Save text to file
        text_filename = f"ocr_text_{uuid.uuid4().hex[:8]}.txt"
        text_path = os.path.join(output_dir, text_filename)
        with open(text_path, "w", encoding="utf-8") as f:
            f.write(combined_text)

        return {
            "success": True,
            "text": combined_text,
            "word_count": len([w for w in combined_text.split() if w.strip()]),
            "char_count": len(combined_text),
            "page_count": page_count,
            "confidence": round(avg_confidence, 1),
            "output_path": text_path,
            "filename": text_filename,
        }

    except Exception as e:
        self.retry(exc=e)


@shared_task(bind=True, max_retries=2, default_retry_delay=60)
def ocr_to_word_task(
    self,
    file_path: str,
    file_type: str,  # 'image' or 'pdf'
    language: str = "eng",
    output_dir: str = None,
) -> dict:
    """Convert scanned image/PDF to Word document.

    Args:
        file_path: Path to the input file
        file_type: 'image' or 'pdf'
        language: Tesseract language code
        output_dir: Directory to save output file

    Returns:
        dict with output file path and metadata
    """
    try:
        from docx import Document
        from docx.shared import Pt

        output_dir = output_dir or settings.temp_file_dir

        # Extract text based on file type
        if file_type == "image":
            result = extract_text_from_image_task(file_path, language, output_dir)
        else:
            result = extract_text_from_pdf_task(file_path, language, output_dir)

        if not result.get("success"):
            return result

        text = result["text"]
        if not text:
            return {"success": False, "error": "No text detected in the file"}

        # Create Word document
        doc = Document()
        doc.add_heading("OCR Extracted Text", 0)
        doc.add_paragraph(
            f"Words: {result['word_count']} | Confidence: {result['confidence']}%"
        )
        doc.add_paragraph("")

        for paragraph in text.split("\n\n"):
            if paragraph.strip():
                p = doc.add_paragraph(paragraph.strip())
                p.style.font.size = Pt(11)

        # Save document
        output_filename = f"ocr_document_{uuid.uuid4().hex[:8]}.docx"
        output_path = os.path.join(output_dir, output_filename)
        doc.save(output_path)

        return {
            "success": True,
            "output_path": output_path,
            "filename": output_filename,
            "size": os.path.getsize(output_path),
            "word_count": result["word_count"],
            "confidence": result["confidence"],
        }

    except Exception as e:
        self.retry(exc=e)
